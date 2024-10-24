#Dataframe to Docx

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import matplotlib.pyplot as plt
# Sample DataFrames (replace these with your actual DataFrames)
df1 = pd.DataFrame({'A': [1, 2, 3], 'B': [4, 5, 6]})
df2 = pd.DataFrame({'X': ['a', 'b', 'c'], 'Y': ['d', 'e', 'f']})

# Create a Word document
doc = Document()
# Sample Matplotlib chart
plt.figure(figsize=(6, 4))
plt.plot([1, 2, 3], [4, 5, 6], label='Line 1')
plt.plot([1, 2, 3], [6, 5, 4], label='Line 2')
plt.xlabel('X-axis')
plt.ylabel('Y-axis')
plt.legend()
plt.title('Sample Chart')

# Save the chart as an image
chart_image_path = 'chart_image.png'
plt.savefig(chart_image_path, format='png')
plt.close()

# Create a Word document
doc = Document()

# Add a title to the document
doc.add_heading('DataFrames and Chart in Word Document', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



# Add a heading for the chart
doc.add_heading('Chart Explanation', level=2)

# Add the chart image to the document
doc.add_picture(chart_image_path, width=Inches(6))
# Add a title to the document
doc.add_heading('DataFrames in Word Document', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Function to add a DataFrame to the document
def add_dataframe_to_doc(df, doc, title):
    # Add a heading for the DataFrame
    doc.add_heading(title, level=2) 

    # Convert DataFrame to a table
    table = doc.add_table(df.shape[0] + 1, df.shape[1])
    table.autofit = True

    # Add column names to the table
    for col_num, col_name in enumerate(df.columns):
        table.cell(0, col_num).text = col_name

    # Add data to the table
    for row_num in range(df.shape[0]):
        for col_num in range(df.shape[1]):
            table.cell(row_num + 1, col_num).text = str(df.iloc[row_num, col_num])
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Some explanation of the DataFrame goes here.')

# Add DataFrames to the document
add_dataframe_to_doc(df1, doc, 'DataFrame 1 Explanation')
add_dataframe_to_doc(df2, doc, 'DataFrame 2 Explanation')

# Save the document
doc.save('output_document.docx')
